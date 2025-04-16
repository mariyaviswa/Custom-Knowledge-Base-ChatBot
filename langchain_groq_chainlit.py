# --- Environment Setup ---
from dotenv import load_dotenv
import os

# Load environment variables from .env file
load_dotenv()

# Get API keys
GROQ_API_KEY = os.getenv("GROQ_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
LANGCHAIN_API_KEY = os.getenv("LANGCHAIN_API_KEY")

if not GROQ_API_KEY and not OPENAI_API_KEY and not LANGCHAIN_API_KEY:
    raise ValueError("Please set at least one of GROQ_API_KEY, OPENAI_API_KEY, or LANGCHAIN_API_KEY in your .env file.")

# --- Imports ---
from langchain_groq import ChatGroq
from langchain.chat_models import ChatOpenAI
from langchain_community.chat_models import ChatAnthropic
from langchain_core.prompts import ChatPromptTemplate
from langchain.schema import StrOutputParser
from langchain.schema.runnable import Runnable
from langchain.schema.runnable.config import RunnableConfig
from langchain.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
import chainlit as cl
import traceback
import pandas as pd
from docx import Document

# File processing helpers
async def process_csv(file_path: str) -> str:
    try:
        df = pd.read_csv(file_path)
        return df.to_string(index=False)
    except Exception as e:
        return f"Error processing CSV file: {e}"

async def process_excel(file_path: str) -> str:
    try:
        df = pd.read_excel(file_path)
        return df.to_string(index=False)
    except Exception as e:
        return f"Error processing Excel file: {e}"

async def process_word(file_path: str) -> str:
    try:
        doc = Document(file_path)
        if not doc.paragraphs:
            return "The Word document is empty."
        content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return content if content else "The Word document contains no readable text."
    except Exception as e:
        return f"Error processing Word file: {e}"

async def process_file(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        loader = PyPDFLoader(file_path)
        pages = loader.load()
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        docs = splitter.split_documents(pages)
        return "\n\n".join([doc.page_content for doc in docs])
    elif ext == ".csv":
        return await process_csv(file_path)
    elif ext in [".xls", ".xlsx"]:
        return await process_excel(file_path)
    elif ext in [".doc", ".docx"]:
        return await process_word(file_path)
    else:
        return "Unsupported file type. Please upload a PDF, CSV, Excel, or Word (.docx) file."

# --- Chat Start ---
@cl.on_chat_start
async def on_chat_start():
    await cl.Message(
        content="Hello! My name is Kelsy 🤖.I am a chatbot interact with me using Upload a file or ask any question about the file.",
        elements=[cl.Image(name="image1", display="inline", path="chat.png")]
    ).send()

    backend = os.getenv("LANGCHAIN_BACKEND", "groq").lower()

    if backend == "groq":
        model = ChatGroq(temperature=0, model_name="llama3-70b-8192", groq_api_key=GROQ_API_KEY)
    elif backend == "openai":
        model = ChatOpenAI(temperature=0, model_name="gpt-4", openai_api_key=OPENAI_API_KEY)
    elif backend == "langchain":
        model = ChatAnthropic(temperature=0, model_name="claude-3-opus-20240229", api_key=LANGCHAIN_API_KEY)
    else:
        raise ValueError("Unsupported backend selected.")

    prompt = ChatPromptTemplate.from_messages([
        ("system", "You're a knowledgeable Machine Learning Engineer."),
        ("human", "{question}")
    ])

    cl.user_session.set("runnable", prompt | model | StrOutputParser())
    cl.user_session.set("file_text", "")  # Initialize file content

# --- On User Message ---
@cl.on_message
async def on_message(message: cl.Message):
    runnable = cl.user_session.get("runnable")
    msg = cl.Message(content="")

    stored_file_text = cl.user_session.get("file_text") or ""
    current_file_text = ""

    if message.elements:
        for element in message.elements:
            if isinstance(element, cl.File):
                current_file_text = await process_file(element.path)
                if "Unsupported file type" in current_file_text:
                    await cl.Message(content=current_file_text).send()
                    return
                cl.user_session.set("file_text", current_file_text)
                stored_file_text = current_file_text  # update for current response

    combined_input = f"User Query: {message.content}\n\nFile Content:\n{stored_file_text}" if stored_file_text else message.content

    try:
        async for chunk in runnable.astream(
            {"question": combined_input},
            config=RunnableConfig(callbacks=[cl.LangchainCallbackHandler()])
        ):
            await msg.stream_token(chunk)
    except Exception:
        msg.content = f"An error occurred:\n```\n{traceback.format_exc()}\n```"
        await msg.send()
    else:
        await msg.send()
